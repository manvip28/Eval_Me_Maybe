import json
import time
import os
from pathlib import Path
from .utils.pdf_extractor import extract_topics as extract_topics_pdf
from .utils.docx_extractor import extract_topics as extract_topics_docx, extract_topics_with_images
from .utils.keyword_extractor import KeywordExtractor
from .question_generation.question_generator import QuestionGenerator
from .utils.manual_review import review_questions
from .question_generation.context_generator import ContentGenerator
from .utils.image_associator import ImageAssociator
from docx import Document
from docx.shared import Pt, Inches
from docx.table import Table
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

def save_to_docx(questions, docx_file="final_questions.docx", image_associator=None):
    """Save approved questions to a Word file with clear formatting and bold labels."""
    doc = Document()
    doc.add_heading("Final Approved Questions", level=1)

    for i, q in enumerate(questions, start=1):
        # Add question number
        doc.add_paragraph(f"Q{i}.")

        # Add each field with bold label
        p = doc.add_paragraph()
        p.add_run("Question: ").bold = True
        p.add_run(q.get("question", ""))

        # Add diagram if available (can be string or list)
        if "diagram" in q and q["diagram"]:
            try:
                # Add image to document
                doc.add_paragraph("Diagram:")
                
                # Handle both string and list cases
                diagram_refs = q["diagram"]
                if isinstance(diagram_refs, str):
                    diagram_refs = [diagram_refs]
                elif not isinstance(diagram_refs, list):
                    diagram_refs = []
                
                # Load image from Azure blob storage
                import tempfile
                from storage import get_storage_client
                
                storage = get_storage_client()
                images_added = 0
                
                for diagram_ref in diagram_refs:
                    if not diagram_ref:
                        continue
                    
                    # Convert relative path to blob path (remove ./ if present)
                    blob_path = diagram_ref.lstrip('./') if isinstance(diagram_ref, str) and diagram_ref.startswith('./') else diagram_ref
                    
                    if storage.is_blob_storage() and storage.exists(blob_path):
                        try:
                            # Download image from blob storage to temp file
                            image_data = storage.read_file(blob_path)
                            
                            # Create temporary file for the image
                            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_img:
                                tmp_img_path = tmp_img.name
                                tmp_img.write(image_data)
                            
                            try:
                                # Add image to document
                                doc.add_picture(tmp_img_path, width=Inches(4.0))
                                images_added += 1
                            finally:
                                # Clean up temp file
                                try:
                                    os.unlink(tmp_img_path)
                                except:
                                    pass
                        except Exception as img_error:
                            print(f"Warning: Could not add image {blob_path}: {img_error}")
                    else:
                        print(f"Warning: Diagram not found in Azure blob storage: {blob_path}")
                
                if images_added == 0:
                    # If no images were added, show the reference
                    doc.add_paragraph(f"[Image: {q['diagram']}]")
                else:
                    doc.add_paragraph("")  # Add space after image(s)
            except Exception as e:
                print(f"Warning: Could not add image {q.get('diagram', 'unknown')}: {e}")
                import traceback
                traceback.print_exc()
                doc.add_paragraph(f"[Image: {q.get('diagram', 'unknown')}]")

        p = doc.add_paragraph()
        p.add_run("Bloom: ").bold = True
        p.add_run(q.get("bloom", ""))

        p = doc.add_paragraph()
        p.add_run("Marks: ").bold = True
        p.add_run(str(q.get("marks", "")))

        keywords = ", ".join(q.get("keywords_used", []))
        p = doc.add_paragraph()
        p.add_run("Keywords Used: ").bold = True
        p.add_run(keywords)

        # Add answer if available
        if "answer" in q:
            p = doc.add_paragraph()
            p.add_run("Answer: ").bold = True
            p.add_run(q.get("answer", ""))

        # Add a blank line between questions
        doc.add_paragraph("")

    # Optional: adjust font size for readability
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    doc.save(docx_file)
    print(f"Questions also saved to {docx_file}")

def save_question_paper_to_docx(questions, docx_file="question_paper.docx"):
    """
    Save questions to a Word file in table format for question paper.
    Format: Q1 | Question | Bloom's Level | Marks
    """
    doc = Document()
    
    # Reduce page margins for a cleaner layout
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    
    doc.add_heading("Question Paper", level=1)
    
    # Create table with headers
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Add header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'QNo.'
    hdr_cells[1].text = 'Question'
    hdr_cells[2].text = "Bloom's Level"
    hdr_cells[3].text = 'Marks'
    
    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add question rows
    for i, q in enumerate(questions, start=1):
        row = table.add_row()
        cells = row.cells
        cells[0].text = f"Q{i}"
        cells[1].text = q.get("question", "")
        cells[2].text = q.get("bloom", "")
        cells[3].text = str(q.get("marks", ""))
    
    # Set specific column widths for better space utilization
    # Use a more direct approach to control column widths
    

    # Disable autofit
    table.allow_autofit = False
    table.autofit = False

    # Define column widths in inches
    col_widths = [Inches(0.7), Inches(5.0), Inches(1.0), Inches(0.6)]

    for row in table.rows:
        for i, width in enumerate(col_widths):
            cell = row.cells[i]
            cell.width = width
            
            # Set the width at XML level
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(width.inches * 1440)))  # width in twips
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

    
    # Optional: adjust font size for readability
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    doc.save(docx_file)
    print(f"Question paper saved to {docx_file}")

def run_pipeline_web(input_path: str, output_file: str = "generated_questions.json", questions_per_topic: int = None):
    """Pipeline execution for web interface - always skips manual review"""
    print("🌐 WEB INTERFACE MODE - Starting pipeline for Streamlit")
    print(f"DEBUG: run_pipeline_web called with skip_manual_review=True")
    print(f"DEBUG: About to call run_pipeline with skip_manual_review=True")
    return run_pipeline(input_path, output_file, questions_per_topic, skip_manual_review=True)

def run_pipeline(input_path: str, output_file: str = "generated_questions.json", questions_per_topic: int = None, skip_manual_review: bool = False):
    """Main pipeline execution"""
    print("Starting textbook processing pipeline")
    print(f"DEBUG: skip_manual_review parameter = {skip_manual_review}")
    print(f"DEBUG: type of skip_manual_review = {type(skip_manual_review)}")
    
    # 1. Text Extraction
    print("Extracting topics from textbook...")
    start = time.time()
    
    # Use appropriate extractor based on file type
    topic_images = {}  # Store images associated with topics
    if input_path.lower().endswith('.pdf'):
        topics = extract_topics_pdf(input_path)
        print(f"Extracted {len(topics)} topics from PDF in {time.time()-start:.1f}s")
    elif input_path.lower().endswith('.docx'):
        # Extract both topics and images together
        topics, topic_images = extract_topics_with_images(input_path, "answer_key_gen")
        print(f"Extracted {len(topics)} topics from DOCX in {time.time()-start:.1f}s")
        print(f"Extracted images for {len([t for t in topic_images if topic_images[t]])} topics")
    else:
        raise ValueError(f"Unsupported file type: {input_path}. Only PDF and DOCX files are supported.")
    
    # 2. Keyword Extraction
    print("Extracting keywords...")
    keyword_extractor = KeywordExtractor()
    start = time.time()
    keywords_dict = keyword_extractor.process_topics(topics)
    print(f"Extracted keywords in {time.time()-start:.1f}s")
    
    
    # 3. Get questions per topic from user if not provided
    if questions_per_topic is None:
        while True:
            try:
                questions_per_topic = int(input(f"\nHow many questions do you want to generate per topic? (Found {len(topics)} topics): "))
                if questions_per_topic > 0:
                    break
                else:
                    print("Please enter a positive number.")
            except ValueError:
                print("Please enter a valid number.")
    
    print(f"Will generate {questions_per_topic} questions per topic ({len(topics)} topics = {len(topics) * questions_per_topic} total questions)")
    
    # 4. Question Generation
    print("Generating questions...")
    qg = QuestionGenerator()
    start = time.time()
    questions = qg.generate_for_topics(topics, keywords_dict, questions_per_topic=questions_per_topic)
    print(f"Generated {len(questions)} questions in {time.time()-start:.1f}s")
    
    # 4.5. Associate images with questions (topic-based approach)
    if topic_images:
        print("Associating images with questions based on topics...")
        
        for i, question in enumerate(questions):
            question_id = f"Q{i+1}"
            topic = list(topics.keys())[i // questions_per_topic] if i // questions_per_topic < len(topics) else "Unknown"
            
            # Get images associated with this topic
            topic_image_list = topic_images.get(topic, [])
            
            if topic_image_list:
                # Use the first image for this topic
                image_info = topic_image_list[0]
                question["diagram"] = image_info["relative_path"]
                
                # Mark this image as used (remove from available list)
                topic_images[topic].pop(0)
            else:
                question["diagram"] = None
    else:
        print("No images available - skipping image association")
        for question in questions:
            question["diagram"] = None
    
    # 4. Save intermediate results to Azure blob storage ONLY
    intermediate_json_path_str = "answer_key_gen/intermediate_questions.json"
    
    # Save to blob storage ONLY
    try:
        from storage import get_storage_client
        storage = get_storage_client()
        if not storage.is_blob_storage():
            raise RuntimeError("Blob storage is not configured. Please configure Azure storage.")
        
        storage.write_json(intermediate_json_path_str, questions)
        print(f"✅ Saved intermediate questions to Azure blob storage: {intermediate_json_path_str}")
    except Exception as e:
        print(f"❌ Failed to save intermediate questions to blob storage: {e}")
        raise
    
    # 5. Manual Review (skip if called from Streamlit)
    print(f"DEBUG: About to check skip_manual_review = {skip_manual_review}")
    
    # Load questions from Azure blob storage for manual review
    try:
        from storage import get_storage_client
        storage = get_storage_client()
        if not storage.is_blob_storage():
            raise RuntimeError("Blob storage is not configured.")
        
        all_questions = storage.read_json(intermediate_json_path_str)
    except Exception as e:
        print(f"❌ Failed to load intermediate questions from blob storage: {e}")
        raise
    
    if not skip_manual_review:
        print("\nStarting manual review process...")
        print("🔍 CLI Manual Review Mode - Interactive review in terminal")
        # For CLI review, we need a local temp file
        temp_review_path = Path("temp_review_questions.json")
        with open(temp_review_path, "w") as f:
            json.dump(all_questions, f, indent=2)
        review_questions(temp_review_path)
        
        # Reload from temp file after review
        with open(temp_review_path, "r") as f:
            all_questions = json.load(f)
        temp_review_path.unlink()  # Clean up temp file
        
        # 6. Filter and save only approved questions
        print("\nFiltering approved questions...")
        approved_questions = [q for q in all_questions if q.get("approved", False)]
        
        # Save updated questions back to Azure
        storage.write_json(intermediate_json_path_str, all_questions)
    else:
        print("\n🌐 WEB INTERFACE MODE - Skipping CLI manual review")
        print("📝 Questions will be reviewed in the web interface")
        approved_questions = all_questions  # Don't filter, let web interface handle approval
    
    # Save approved questions to Azure blob storage ONLY
    final_json_path_str = "answer_key_gen/final_questions.json"
    final_docx_path_str = "answer_key_gen/final_answer_key.docx"
    question_paper_path_str = "answer_key_gen/question_paper.docx"
    
    # Save final JSON to blob storage ONLY
    try:
        from storage import get_storage_client
        storage = get_storage_client()
        if not storage.is_blob_storage():
            raise RuntimeError("Blob storage is not configured.")
        
        storage.write_json(final_json_path_str, approved_questions)
        print(f"✅ Saved final questions to Azure blob storage: {final_json_path_str}")
    except Exception as e:
        print(f"❌ Failed to save final questions to blob storage: {e}")
        raise

    # Only generate final documents if not skipping manual review
    if not skip_manual_review:
        # Generate to temp files first, then upload to Azure
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx1:
            temp_docx1_path = tmp_docx1.name
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx2:
            temp_docx2_path = tmp_docx2.name
        
        try:
            save_to_docx(approved_questions, temp_docx1_path, None)
            save_question_paper_to_docx(approved_questions, temp_docx2_path)
            
            # Upload final DOCX files to blob storage ONLY
            try:
                from storage import get_storage_client
                storage = get_storage_client()
                if not storage.is_blob_storage():
                    raise RuntimeError("Blob storage is not configured.")
                
                # Upload answer key
                with open(temp_docx1_path, "rb") as f:
                    storage.write_file(final_docx_path_str, f.read())
                print(f"✅ Saved final answer key to Azure blob storage: {final_docx_path_str}")
                
                # Upload question paper
                with open(temp_docx2_path, "rb") as f:
                    storage.write_file(question_paper_path_str, f.read())
                print(f"✅ Saved question paper to Azure blob storage: {question_paper_path_str}")
            except Exception as e:
                print(f"❌ Failed to upload final documents to blob storage: {e}")
                raise
        finally:
            # Clean up temp files
            try:
                os.unlink(temp_docx1_path)
                os.unlink(temp_docx2_path)
            except:
                pass
        
        print(f"\n✅ Pipeline completed!")
        print(f"Total questions generated: {len(all_questions)}")
        print(f"Approved questions: {len(approved_questions)}")
        print(f"Intermediate questions saved to Azure: {intermediate_json_path_str}")
        print(f"Final questions saved to Azure: {final_json_path_str}")
        print(f"Final answer key document saved to Azure: {final_docx_path_str}")
        print(f"Question paper saved to Azure: {question_paper_path_str}")
        print(f"Diagrams saved to Azure: answer_key_gen/diagrams/")
    else:
        print(f"\n🎉 Pipeline completed (WEB INTERFACE MODE)!")
        print(f"📊 Total questions generated: {len(all_questions)}")
        print(f"💾 Intermediate questions saved to Azure: {intermediate_json_path_str}")
        print(f"🌐 Manual review will be done in the web interface")
        print(f"📄 Final documents will be generated after web-based review")
        print(f"✅ Ready for Streamlit web interface!")




if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python pipeline.py <textbook_path.pdf> [output_file.json]")
        sys.exit(1)
    
    output_file = sys.argv[2] if len(sys.argv) > 2 else "generated_questions.json"
    
    # Check for skip-review flag
    skip_manual_review = "--skip-review" in sys.argv
    
    run_pipeline(sys.argv[1], output_file, skip_manual_review=skip_manual_review)