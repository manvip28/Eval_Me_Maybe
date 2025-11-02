"""
DOCX Text Extraction Module
Extracts topic-wise text from DOCX files using headings
"""

import re
import os
from collections import OrderedDict
from typing import OrderedDict as OrderedDictType, List, Dict, Optional
from docx import Document
from pathlib import Path


def is_topic_header(text: str) -> bool:
    """Return True if the line looks like a topic header."""
    text = text.strip()
    
    # Pattern 1: Numbered headers like "1.1 Title"
    numbered_pattern = r"^\d+(\.\d+)*\s+.+$"
    if re.match(numbered_pattern, text):
        return True
    
    # Pattern 2: Topic-based headers (all caps, short phrases)
    # Examples: "INTRODUCTION", "WELL-POSED LEARNING PROBLEMS", "CONCEPT LEARNING"
    if (text.isupper() and 
        len(text.split()) <= 6 and  # Not too long
        len(text) > 3 and  # Not too short
        not text.endswith('.') and  # Not a sentence
        not '=' in text and  # Not a mathematical expression
        not text.startswith('Fig:') and  # Not a figure caption
        not text.startswith('Table:') and  # Not a table caption
        not re.match(r'^\d+$', text)):  # Not just a number
        return True
    
    return False


def clean_header(text: str) -> str:
    """
    Cleans section headers.
    
    For numbered headers: '1.1.   Cloud Computing' → '1.1 Cloud Computing'
    For topic headers: 'INTRODUCTION' → 'INTRODUCTION'
    """
    text = text.strip()

    # Handle numbered headers
    if re.match(r"^\d+(?:\.\d+)*\.\s", text):
        # Fix common issue: '1.1.' → '1.1'
        text = re.sub(r"^(\d+(?:\.\d+)*)(\.)\s+", r"\1 ", text)
    else:
        text = re.sub(r"\s+", " ", text)  # remove weird extra spaces

    # For numbered headers, parse section number and title
    numbered_match = re.match(r"^(\d+(?:\.\d+)*)(?:\.)?\s+(.+)$", text)
    if numbered_match:
        section_number, title = numbered_match.groups()
        return f"{section_number} {title.strip()}"
    
    # For topic headers, just return the cleaned text
    return text


def extract_topics(docx_path: str) -> OrderedDictType[str, str]:
    """
    Extract topic-wise text from a DOCX file.

    Identifies headers using is_topic_header, cleans them, and aggregates
    subsequent lines into the current topic until the next header.
    Returns an OrderedDict to preserve section order.
    """
    content: OrderedDictType[str, str] = OrderedDict()
    current_topic: str = ""

    doc = Document(docx_path)
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        
        if is_topic_header(text):
            current_topic = clean_header(text)
            if current_topic not in content:
                content[current_topic] = ""
        elif current_topic:
            # Accumulate text under current topic
            content[current_topic] += text + "\n"

    return content


def extract_topics_with_images(docx_path: str, output_dir: str = "answer_key_gen") -> tuple[OrderedDictType[str, str], Dict[str, List[Dict]]]:
    """
    Extract topic-wise text AND images from a DOCX, associating images with topics.
    
    Returns:
        - topics: OrderedDict of topic -> content
        - topic_images: Dict of topic -> list of associated images
    """
    content: OrderedDictType[str, str] = OrderedDict()
    topic_images: Dict[str, List[Dict]] = {}
    current_topic: str = ""
    
    # Create output directory for images
    diagrams_dir = Path(output_dir) / "diagrams"
    diagrams_dir.mkdir(parents=True, exist_ok=True)
    
    doc = Document(docx_path)
    
    # First pass: extract text and identify topics
    header_count = 0
    
    # Extract all text from DOCX, including headings
    all_text = ""
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            # Check if this is a heading (style-based detection)
            style_name = paragraph.style.name if paragraph.style else "Normal"
            is_heading = "Heading" in style_name or "Title" in style_name
            
            if is_heading:
                # Add heading marker to make it easier to detect
                all_text += f"HEADING: {text}\n"
            else:
                all_text += text + "\n"
    
    # Process line by line
    lines = all_text.split("\n")
    
    for line in lines:
        # Check for heading markers first
        if line.startswith("HEADING:"):
            heading_text = line.replace("HEADING:", "").strip()
            if is_topic_header(heading_text):
                current_topic = clean_header(heading_text)
                if current_topic not in content:
                    content[current_topic] = ""
                    topic_images[current_topic] = []
                    header_count += 1
        elif is_topic_header(line):
            current_topic = clean_header(line)
            if current_topic not in content:
                content[current_topic] = ""
                topic_images[current_topic] = []
                header_count += 1
        elif current_topic:
            # Accumulate text under current topic
            content[current_topic] += line + "\n"
    
    # Second pass: extract images and associate with proper topics
    image_count = 0
    
    # Filter out mathematical expressions and get proper topics
    proper_topics = [t for t in content.keys() 
                    if not '=' in t and 
                    not t.startswith('3.2.2') and
                    not t.startswith('5.4.4') and
                    not t.startswith('1 +') and
                    not t.startswith('V :') and
                    not t.startswith('FIND-') and
                    len(t) > 5]  # Not too short
    
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                # Get image data
                image_data = rel.target_part.blob
                
                # Check image size - skip very small images (likely symbols/icons)
                if len(image_data) < 20000:  # Skip images smaller than 20KB
                    continue
                
                # Try to get actual image dimensions to filter out text snippets
                try:
                    from PIL import Image
                    import io
                    img_obj = Image.open(io.BytesIO(image_data))
                    width, height = img_obj.size
                    
                    # Skip images that are too small (likely text snippets or icons)
                    if width < 200 or height < 200:
                        continue
                        
                    # Update size with actual dimensions
                    actual_size = (width, height)
                    img_obj.close()
                except:
                    # If we can't get dimensions, use default size
                    actual_size = (200, 200)
                
                # Save image with topic-based naming
                img_filename = f"diagram_{image_count}.png"
                img_path = diagrams_dir / img_filename
                
                with open(img_path, 'wb') as f:
                    f.write(image_data)
                
                # Associate with proper topics in round-robin fashion
                if proper_topics:
                    associated_topic = proper_topics[image_count % len(proper_topics)]
                else:
                    associated_topic = "Unknown"
                
                # Store image metadata
                image_info = {
                    "path": str(img_path),
                    "filename": img_filename,
                    "relative_path": f"./answer_key_gen/diagrams/{img_filename}",
                    "size": actual_size,
                    "topic": associated_topic
                }
                
                topic_images[associated_topic].append(image_info)
                image_count += 1
                
            except Exception as e:
                print(f"Warning: Could not extract image: {e}")
                continue
    
    return content, topic_images


__all__ = [
    "is_topic_header",
    "clean_header", 
    "extract_topics",
    "extract_topics_with_images",
]
