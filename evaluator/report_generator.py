#!/usr/bin/env python3
"""
Generate human-readable feedback reports from evaluation_results.json
"""

import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_question_feedback(q_num, result):
    """Generate detailed feedback for a single question."""
    question = result.get("question", "N/A")
    student_answer = result.get("student_answer", "No answer provided")
    expected_answer = result.get("expected_answer", "No reference answer")
    percentage_score = result.get("percentage_score", 0)

    # Include image info if available
    has_student_image = result.get("has_student_image", False)
    has_reference_image = result.get("has_reference_image", False)
    image_similarity = None
    if "evaluation_details" in result:
        image_similarity = result["evaluation_details"].get("image_similarity")

    feedback = f"""
## Question {q_num}
**Question:** {question}

**Your Answer:** {student_answer}
**Reference Answer:** {expected_answer}

**Score:** {percentage_score}%
"""
    if has_student_image or has_reference_image:
        feedback += f"**Image Provided:** Student: {has_student_image}, Reference: {has_reference_image}\n"
        if image_similarity is not None:
            feedback += f"**Image Similarity:** {image_similarity}\n"

    return feedback + "\n---\n"

def generate_summary(results):
    """Print console summary based on individual_results and summary."""
    individual_results = results.get("individual_results", {})
    summary = results.get("summary", {})

    print("=" * 50)
    print("           EVALUATION SUMMARY")
    print("=" * 50)

    total_score = 0
    total_questions = len(individual_results)

    for q_num, res in individual_results.items():
        score = res.get("percentage_score", 0)
        print(f"{q_num}: {score}%")
        total_score += score

    avg_score = total_score / total_questions if total_questions else 0
    print("-" * 50)
    print(f"Average Score: {avg_score:.1f}%")

    # Simple overall rating
    if avg_score >= 90:
        overall = "Excellent!"
    elif avg_score >= 80:
        overall = "Very Good!"
    elif avg_score >= 70:
        overall = "Good"
    elif avg_score >= 60:
        overall = "Satisfactory"
    else:
        overall = "Needs Improvement"
    print(f"Overall: {overall}")
    print("=" * 50)

def generate_full_report(results_json_path, output_path="feedback_report.md"):
    """Generate full Markdown report for all questions."""
    with open(results_json_path, 'r', encoding='utf-8') as f:
        results = json.load(f)

    individual_results = results.get("individual_results", {})
    summary = results.get("summary", {})

    report = f"# Student Answer Evaluation Report\n"
    report += f"**Date:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
    report += f"## Overall Summary\n"
    report += f"- Total Questions: {summary.get('total_questions', len(individual_results))}\n"
    report += f"- Questions Answered: {summary.get('answered_questions', 0)}\n"
    report += f"- Questions Evaluated: {summary.get('evaluated_questions', 0)}\n"
    report += f"- Overall Average Score: {summary.get('overall_average', 0)}%\n"
    report += f"- Total Achieved Score: {summary.get('total_achieved_score', 0)}\n"
    report += f"- Total Possible Score: {summary.get('total_possible_score', 0)}\n\n"

    report += "## Detailed Question Feedback\n"

    for q_num, res in individual_results.items():
        report += generate_question_feedback(q_num, res)

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(report)

    print(f"\n📄 Detailed feedback report saved to: {output_path}")
    return results

def generate_docx_report(results_json_path, output_path="evaluation_report.docx"):
    """Generate a formatted DOCX report for all questions."""
    with open(results_json_path, 'r', encoding='utf-8') as f:
        results = json.load(f)

    individual_results = results.get("individual_results", {})
    summary = results.get("summary", {})
    
    # Create document
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading("Student Answer Evaluation Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph("")  # Blank line
    
    # Overall Summary Section
    doc.add_heading("Overall Summary", level=2)
    
    # Create summary table
    summary_table = doc.add_table(rows=7, cols=2)
    summary_table.style = 'Light Grid Accent 1'
    
    summary_data = [
        ("Total Questions", str(summary.get('total_questions', len(individual_results)))),
        ("Questions Answered", str(summary.get('answered_questions', 0))),
        ("Questions Evaluated", str(summary.get('evaluated_questions', 0))),
        ("Overall Average Score", f"{summary.get('overall_average', 0)}%"),
        ("Total Achieved Score", str(summary.get('total_achieved_score', 0))),
        ("Total Possible Score", str(summary.get('total_possible_score', 0))),
        ("Overall Rating", get_overall_rating(summary.get('overall_average', 0)))
    ]
    
    for i, (label, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value
        
        # Make label bold
        for paragraph in summary_table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)
        
        # Set value font
        for paragraph in summary_table.rows[i].cells[1].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)
    
    doc.add_paragraph("")  # Blank line
    
    # Detailed Question Feedback Section
    doc.add_heading("Detailed Question Feedback", level=2)
    
    # Add each question
    for q_num, result in individual_results.items():
        # Question heading
        doc.add_heading(f"Question {q_num}", level=3)
        
        # Question text
        p = doc.add_paragraph()
        p.add_run("Question: ").bold = True
        p.add_run(result.get("question", "N/A"))
        
        # Student Answer
        p = doc.add_paragraph()
        p.add_run("Your Answer: ").bold = True
        student_answer = result.get("student_answer", "No answer provided")
        p.add_run(student_answer)
        
        # Reference Answer
        p = doc.add_paragraph()
        p.add_run("Reference Answer: ").bold = True
        expected_answer = result.get("expected_answer", "No reference answer")
        p.add_run(expected_answer)
        
        # Score
        p = doc.add_paragraph()
        score_run = p.add_run("Score: ").bold = True
        percentage_score = result.get("percentage_score", 0)
        score_value = p.add_run(f"{percentage_score}%")
        
        # Color code the score
        if percentage_score >= 80:
            score_value.font.color.rgb = RGBColor(0, 128, 0)  # Green
        elif percentage_score >= 60:
            score_value.font.color.rgb = RGBColor(255, 165, 0)  # Orange
        else:
            score_value.font.color.rgb = RGBColor(255, 0, 0)  # Red
        
        # Image info if available
        has_student_image = result.get("has_student_image", False)
        has_reference_image = result.get("has_reference_image", False)
        image_similarity = None
        if "evaluation_details" in result:
            image_similarity = result["evaluation_details"].get("image_similarity")
        
        if has_student_image or has_reference_image:
            p = doc.add_paragraph()
            p.add_run("Image Provided: ").bold = True
            p.add_run(f"Student: {has_student_image}, Reference: {has_reference_image}")
            
            if image_similarity is not None:
                p = doc.add_paragraph()
                p.add_run("Image Similarity: ").bold = True
                p.add_run(f"{image_similarity:.2f}")
        
        # Evaluation details if available
        if "evaluation_details" in result:
            details = result["evaluation_details"]
            p = doc.add_paragraph()
            p.add_run("Evaluation Metrics: ").bold = True
            
            metrics = []
            if details.get("semantic_score") is not None:
                metrics.append(f"Semantic Score: {details.get('semantic_score', 0):.2f}")
            if details.get("bleu") is not None:
                metrics.append(f"BLEU Score: {details.get('bleu', 0):.2f}")
            if details.get("rouge_l") is not None:
                metrics.append(f"ROUGE-L Score: {details.get('rouge_l', 0):.2f}")
            
            if metrics:
                p.add_run(" | ".join(metrics))
        
        # Add separator line
        doc.add_paragraph("─" * 50)
        doc.add_paragraph("")  # Blank line
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Save document
    doc.save(output_path)
    print(f"\n📄 Detailed DOCX feedback report saved to: {output_path}")
    return results

def get_overall_rating(avg_score):
    """Get overall rating based on average score."""
    if avg_score >= 90:
        return "Excellent!"
    elif avg_score >= 80:
        return "Very Good!"
    elif avg_score >= 70:
        return "Good"
    elif avg_score >= 60:
        return "Satisfactory"
    else:
        return "Needs Improvement"

def generate_multi_student_docx_report(all_results, output_path="evaluation_report_all_students.docx"):
    """Generate a formatted DOCX report for multiple students in one document."""
    # Create document
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading("Student Answer Evaluation Report - All Students", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph("")  # Blank line
    
    # Process each student
    for student_idx, results in enumerate(all_results, start=1):
        individual_results = results.get("individual_results", {})
        summary = results.get("summary", {})
        student_name = results.get("student_name", f"Student {student_idx}")
        student_filename = results.get("student_filename", "Unknown")
        
        # Student heading
        student_heading = doc.add_heading(f"Student {student_idx}", level=1)
        student_heading.style.font.size = Pt(16)
        
        # Student name (from filename)
        name_para = doc.add_paragraph()
        name_label_run = name_para.add_run("Name: ")
        name_label_run.bold = True
        name_label_run.font.size = Pt(12)
        name_value_run = name_para.add_run(student_name)
        name_value_run.font.size = Pt(12)
        
        doc.add_paragraph("")  # Blank line
        
        # Overall Summary Section for this student
        doc.add_heading("Overall Summary", level=2)
        
        # Create summary table
        summary_table = doc.add_table(rows=7, cols=2)
        summary_table.style = 'Light Grid Accent 1'
        
        summary_data = [
            ("Total Questions", str(summary.get('total_questions', len(individual_results)))),
            ("Questions Answered", str(summary.get('answered_questions', 0))),
            ("Questions Evaluated", str(summary.get('evaluated_questions', 0))),
            ("Overall Average Score", f"{summary.get('overall_average', 0)}%"),
            ("Total Achieved Score", str(summary.get('total_achieved_score', 0))),
            ("Total Possible Score", str(summary.get('total_possible_score', 0))),
            ("Overall Rating", get_overall_rating(summary.get('overall_average', 0)))
        ]
        
        for i, (label, value) in enumerate(summary_data):
            summary_table.rows[i].cells[0].text = label
            summary_table.rows[i].cells[1].text = value
            
            # Make label bold
            for paragraph in summary_table.rows[i].cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(11)
            
            # Set value font
            for paragraph in summary_table.rows[i].cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
        
        doc.add_paragraph("")  # Blank line
        
        # Detailed Question Feedback Section
        doc.add_heading("Detailed Question Feedback", level=2)
        
        # Add each question
        for q_num, result in individual_results.items():
            # Question heading
            doc.add_heading(f"Question {q_num}", level=3)
            
            # Question text
            p = doc.add_paragraph()
            p.add_run("Question: ").bold = True
            p.add_run(result.get("question", "N/A"))
            
            # Student Answer
            p = doc.add_paragraph()
            p.add_run("Your Answer: ").bold = True
            student_answer = result.get("student_answer", "No answer provided")
            p.add_run(student_answer)
            
            # Reference Answer
            p = doc.add_paragraph()
            p.add_run("Reference Answer: ").bold = True
            expected_answer = result.get("expected_answer", "No reference answer")
            p.add_run(expected_answer)
            
            # Score
            p = doc.add_paragraph()
            score_run = p.add_run("Score: ").bold = True
            percentage_score = result.get("percentage_score", 0)
            score_value = p.add_run(f"{percentage_score}%")
            
            # Color code the score
            if percentage_score >= 80:
                score_value.font.color.rgb = RGBColor(0, 128, 0)  # Green
            elif percentage_score >= 60:
                score_value.font.color.rgb = RGBColor(255, 165, 0)  # Orange
            else:
                score_value.font.color.rgb = RGBColor(255, 0, 0)  # Red
            
            # Image info if available
            has_student_image = result.get("has_student_image", False)
            has_reference_image = result.get("has_reference_image", False)
            image_similarity = None
            if "evaluation_details" in result:
                image_similarity = result["evaluation_details"].get("image_similarity")
            
            if has_student_image or has_reference_image:
                p = doc.add_paragraph()
                p.add_run("Image Provided: ").bold = True
                p.add_run(f"Student: {has_student_image}, Reference: {has_reference_image}")
                
                if image_similarity is not None:
                    p = doc.add_paragraph()
                    p.add_run("Image Similarity: ").bold = True
                    p.add_run(f"{image_similarity:.2f}")
            
            # Evaluation details if available
            if "evaluation_details" in result:
                details = result["evaluation_details"]
                p = doc.add_paragraph()
                p.add_run("Evaluation Metrics: ").bold = True
                
                metrics = []
                if details.get("semantic_score") is not None:
                    metrics.append(f"Semantic Score: {details.get('semantic_score', 0):.2f}")
                if details.get("bleu") is not None:
                    metrics.append(f"BLEU Score: {details.get('bleu', 0):.2f}")
                if details.get("rouge_l") is not None:
                    metrics.append(f"ROUGE-L Score: {details.get('rouge_l', 0):.2f}")
                
                if metrics:
                    p.add_run(" | ".join(metrics))
            
            # Add separator line
            doc.add_paragraph("─" * 50)
            doc.add_paragraph("")  # Blank line
        
        # Add page break between students (except for the last one)
        if student_idx < len(all_results):
            doc.add_page_break()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Save document
    doc.save(output_path)
    print(f"\n📄 Multi-student DOCX feedback report saved to: {output_path}")
    return all_results

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Generate feedback report from evaluation_results.json")
    parser.add_argument("--results", required=True, help="Path to evaluation_results.json")
    parser.add_argument("--output", default="feedback_report.md", help="Output Markdown report file")
    parser.add_argument("--summary-only", action="store_true", help="Only print console summary")
    args = parser.parse_args()

    results = generate_full_report(args.results, args.output) if not args.summary_only else None

    if args.summary_only:
        with open(args.results, 'r', encoding='utf-8') as f:
            results = json.load(f)
    generate_summary(results)
