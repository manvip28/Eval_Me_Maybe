import os
import re
import json
from docx import Document
from docx.oxml.ns import qn

def extract_images_from_cell(cell, doc, qid, questions):
    """
    Extract all inline images from a docx table cell and save them.
    Updates questions[q_num]["diagram"] with image paths.
    """
    # Resolve filesystem and JSON-visible paths for diagrams
    base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    diagrams_dir = os.path.join(base_dir, "output", "diagrams")
    os.makedirs(diagrams_dir, exist_ok=True)

    for drawing in cell._tc.iterchildren():  # iterate over direct child elements
        for elem in drawing.iter():          # iterate recursively
            if elem.tag.endswith('}blip'):
                rId = elem.get(qn('r:embed'))
                if not rId:
                    continue
                image_part = doc.part.related_parts[rId]

                image_index = len(questions[qid]["diagram"]) + 1
                file_name = f"{qid}_diagram_{image_index}.png"
                image_filename_fs = os.path.join(diagrams_dir, file_name)
                with open(image_filename_fs, "wb") as f:
                    f.write(image_part.blob)
                # Store a project-root relative path in JSON
                image_json_path = os.path.join("./output/diagrams", file_name)
                questions[qid]["diagram"].append(image_json_path)


def extract_images_from_paragraph(paragraph, doc, qid, questions):
    """
    Extract all inline images from a docx paragraph and save them.
    Updates questions[qid]["diagram"] with image paths.
    """
    # Resolve filesystem and JSON-visible paths for diagrams
    base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    diagrams_dir = os.path.join(base_dir, "output", "diagrams")
    os.makedirs(diagrams_dir, exist_ok=True)

    # Initialize diagram list if it doesn't exist
    if qid not in questions:
        questions[qid] = {"text": "", "keywords": "", "diagram": []}
    
    if "diagram" not in questions[qid]:
        questions[qid]["diagram"] = []

    # Look for images in the paragraph's runs
    for run in paragraph.runs:
        for drawing in run._element.iter():
            if drawing.tag.endswith('}blip'):
                rId = drawing.get(qn('r:embed'))
                if not rId:
                    continue
                try:
                    image_part = doc.part.related_parts[rId]

                    image_index = len(questions[qid]["diagram"]) + 1
                    file_name = f"{qid}_diagram_{image_index}.png"
                    image_filename_fs = os.path.join(diagrams_dir, file_name)
                    with open(image_filename_fs, "wb") as f:
                        f.write(image_part.blob)
                    # Store a project-root relative path in JSON
                    image_json_path = os.path.join("./output/diagrams", file_name)
                    questions[qid]["diagram"].append(image_json_path)
                except Exception as e:
                    print(f"Warning: Could not extract image from paragraph: {e}")
                    continue


def extract_from_docx_table(docx_file, output_file="../output/final_answer.json"):
    # Ensure output directories exist
    base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    out_dir = os.path.join(base_dir, "output")
    os.makedirs(os.path.join(out_dir, "diagrams"), exist_ok=True)
    os.makedirs(os.path.dirname(os.path.abspath(output_file)), exist_ok=True)

    doc = Document(docx_file)
    questions = {}

    # Regex to match labels like: Q1, Q1:, Q1., Q1- (case-insensitive)
    question_pattern = re.compile(r"^\s*Q\s*(\d+)\s*[:.)-]?\s*$", re.IGNORECASE)
    
    # Regex patterns for the new format
    answer_pattern = re.compile(r"^\s*Answer:\s*(.+)$", re.IGNORECASE)
    keywords_pattern = re.compile(r"^\s*Keywords:\s*(.+)$", re.IGNORECASE)
    diagram_pattern = re.compile(r"^\s*Diagram:\s*$", re.IGNORECASE)

    # Process paragraphs instead of tables for the new format
    current_question = None
    current_answer = ""
    current_keywords = ""
    in_diagram_section = False
    
    
    # Try both paragraph-based and table-based extraction
    # First, try table-based extraction (original method)
    if len(doc.tables) > 0:
        for table in doc.tables:
            for row in table.rows:
                # Handle both single-column and multi-column formats
                if len(row.cells) == 1:
                    # Single column format - parse the content
                    cell_text = row.cells[0].text.strip()
                    if not cell_text:
                        continue
                    
                    # Split by newlines to get individual components
                    lines = cell_text.split('\n')
                    current_qid = None
                    current_answer = ""
                    current_keywords = ""
                    in_diagram_section = False
                    
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                            
                        # Check if this is a question number
                        match = question_pattern.match(line)
                        if match:
                            # Save previous question if exists
                            if current_qid:
                                questions[current_qid] = {
                                    "text": current_answer.strip(),
                                    "keywords": current_keywords.strip(),
                                    "diagram": []
                                }
                            
                            # Start new question
                            q_num = match.group(1)
                            current_qid = f"Q{q_num}"
                            current_answer = ""
                            current_keywords = ""
                            in_diagram_section = False
                            continue
                        
                        # Check for Answer: label
                        answer_match = answer_pattern.match(line)
                        if answer_match and current_qid:
                            current_answer = answer_match.group(1)
                            in_diagram_section = False
                            continue
                            
                        # Check for Keywords: label
                        keywords_match = keywords_pattern.match(line)
                        if keywords_match and current_qid:
                            current_keywords = keywords_match.group(1)
                            in_diagram_section = False
                            continue
                            
                        # Check for Diagram: label
                        if diagram_pattern.match(line) and current_qid:
                            in_diagram_section = True
                            continue
                    
                    # Save the last question
                    if current_qid:
                        questions[current_qid] = {
                            "text": current_answer.strip(),
                            "keywords": current_keywords.strip(),
                            "diagram": []
                        }
                        
                        # Extract images from the cell
                        extract_images_from_cell(row.cells[0], doc, current_qid, questions)
                
                elif len(row.cells) >= 2:
                    # Multi-column format (original method)
                    q_cell = row.cells[0].text.strip()  # First column: Q1, Q2...
                    a_cell = row.cells[1].text.strip()  # Second column: answer text

                    if not q_cell:
                        continue

                    match = question_pattern.match(q_cell)
                    if not match:
                        continue

                    q_num = match.group(1)
                    qid = f"Q{q_num}"
                    question_text = a_cell

                    questions[qid] = {"text": question_text, "keywords": "", "diagram": []}

                    # Extract images from the answer column
                    extract_images_from_cell(row.cells[1], doc, qid, questions)
    
    # If no questions found in tables, try paragraph-based extraction
    if not questions:
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            if not text:
                continue
            
            # Check if this is a question number
            match = question_pattern.match(text)
            if match:
                # Save previous question if exists
                if current_question:
                    questions[current_question] = {
                        "text": current_answer.strip(),
                        "keywords": current_keywords.strip(),
                        "diagram": []
                    }
                
                # Start new question
                q_num = match.group(1)
                current_question = f"Q{q_num}"
                current_answer = ""
                current_keywords = ""
                in_diagram_section = False
                continue
            
            # Check for Answer: label
            answer_match = answer_pattern.match(text)
            if answer_match and current_question:
                current_answer = answer_match.group(1)
                in_diagram_section = False
                continue
                
            # Check for Keywords: label
            keywords_match = keywords_pattern.match(text)
            if keywords_match and current_question:
                current_keywords = keywords_match.group(1)
                in_diagram_section = False
                continue
                
            # Check for Diagram: label
            if diagram_pattern.match(text) and current_question:
                in_diagram_section = True
                continue
                
            # If we're in a diagram section, look for images in this paragraph
            if in_diagram_section and current_question:
                # Extract images from this paragraph
                extract_images_from_paragraph(paragraph, doc, current_question, questions)
                continue
                
            # If we have a current question and this text doesn't match any pattern,
            # it might be continuation of the answer
            if current_question and not in_diagram_section and current_answer:
                current_answer += " " + text

    # Save the last question
    if current_question:
        questions[current_question] = {
            "text": current_answer.strip(),
            "keywords": current_keywords.strip(),
            "diagram": []
        }

    # Clean diagram entries
    for qid, content in questions.items():
        if len(content["diagram"]) == 0:
            content["diagram"] = None
        elif len(content["diagram"]) == 1:
            content["diagram"] = content["diagram"][0]

    # Save JSON
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(questions, f, indent=2, ensure_ascii=False)

    # Avoid Unicode emojis to prevent Windows cp1252 console errors
    print(f"Extracted typed answers + images from {docx_file}")
    print(f"Saved JSON to {output_file}")
    return questions


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python handle_docx.py <answer.docx>")
        exit(1)
    input_file = sys.argv[1]
    extract_from_docx_table(input_file)
